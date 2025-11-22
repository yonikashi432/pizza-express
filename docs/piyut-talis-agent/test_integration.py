#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Integration Test for Piyut-Talis-Agent

This script tests the integration between:
- Piyut-Talis-Agent specification
- Unified Golden Operator seed configuration
- DOCX export functionality
- Text encoding and preservation
"""

import sys
import os
import yaml
import tempfile
from pathlib import Path

# Get the base directory (repository root)
BASE_DIR = Path(__file__).parent.parent.parent

# Test results
results = []

def test_result(name, passed, details=""):
    """Record test result"""
    status = "✓ PASS" if passed else "✗ FAIL"
    results.append((name, passed, details))
    print(f"{status}: {name}")
    if details:
        print(f"  {details}")
    return passed

print("=" * 60)
print("Piyut-Talis-Agent Integration Test")
print("=" * 60)

# Test 1: Agent specification file exists and is valid YAML
print("\n[1] Testing agent specification...")
try:
    spec_path = BASE_DIR / "docs/piyut-talis-agent/agent-specification.yaml"
    with open(spec_path, 'r', encoding='utf-8') as f:
        agent_spec = yaml.safe_load(f)
    
    test_result(
        "Agent specification YAML valid",
        True,
        f"Found {len(agent_spec)} top-level keys"
    )
    
    # Check required fields
    required_fields = ['name', 'capabilities', 'boundaries', 'integration']
    has_all = all(field in agent_spec for field in required_fields)
    test_result(
        "Agent specification has required fields",
        has_all,
        f"Required: {required_fields}"
    )
except Exception as e:
    test_result("Agent specification YAML valid", False, str(e))

# Test 2: Unified seed integration
print("\n[2] Testing unified seed integration...")
try:
    seed_path = BASE_DIR / "seeds/unified-golden-operator-v25-35.yaml"
    with open(seed_path, 'r', encoding='utf-8') as f:
        unified_seed = yaml.safe_load(f)
    
    test_result(
        "Unified seed YAML valid",
        True,
        f"Version: {unified_seed.get('version', 'unknown')}"
    )
    
    # Check agents section exists
    has_agents = 'agents' in unified_seed
    test_result(
        "Agents section exists in unified seed",
        has_agents
    )
    
    if has_agents:
        has_piyut = 'piyut_talis' in unified_seed['agents']
        test_result(
            "Piyut-Talis-Agent registered in unified seed",
            has_piyut,
            f"Agent label: {unified_seed['agents'].get('piyut_talis', {}).get('label', 'N/A')}"
        )
except Exception as e:
    test_result("Unified seed integration", False, str(e))

# Test 3: Canonical text file
print("\n[3] Testing canonical text...")
try:
    text_path = BASE_DIR / "docs/piyut-talis-agent/piyut-canonical.txt"
    with open(text_path, 'r', encoding='utf-8') as f:
        canonical_text = f.read()
    
    # Check encoding
    has_hebrew = any('\u0590' <= c <= '\u05FF' for c in canonical_text)
    test_result(
        "Canonical text contains Hebrew characters",
        has_hebrew,
        f"Text length: {len(canonical_text)} chars, {len(canonical_text.splitlines())} lines"
    )
    
    # Check for nikud (vowel points)
    has_nikud = any('\u0591' <= c <= '\u05C7' for c in canonical_text)
    test_result(
        "Hebrew text includes nikud (vowel points)",
        has_nikud
    )
except Exception as e:
    test_result("Canonical text", False, str(e))

# Test 4: DOCX export functionality
print("\n[4] Testing DOCX export...")
try:
    import_success = False
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import_success = True
    except ImportError:
        pass
    
    test_result(
        "python-docx library available",
        import_success,
        "Required for DOCX export"
    )
    
    if import_success:
        # Import the export function
        agent_dir = str(BASE_DIR / "docs/piyut-talis-agent")
        if agent_dir not in sys.path:
            sys.path.insert(0, agent_dir)
        from export_piyut_docx import create_piyut_docx
        
        # Test export with temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            test_output = tmp_file.name
        
        try:
            result_path = create_piyut_docx(test_output, font_size_pt=16)
            
            file_created = Path(result_path).exists()
            file_size = Path(result_path).stat().st_size if file_created else 0
            
            test_result(
                "DOCX file created successfully",
                file_created,
                f"File size: {file_size:,} bytes"
            )
            
            # Verify DOCX structure
            if file_created:
                doc = Document(result_path)
                has_content = len(doc.paragraphs) > 0
                test_result(
                    "DOCX contains paragraphs",
                    has_content,
                    f"Found {len(doc.paragraphs)} paragraphs"
                )
        finally:
            # Clean up temporary file
            try:
                if Path(test_output).exists():
                    Path(test_output).unlink()
            except (OSError, PermissionError):
                pass  # Ignore cleanup errors
except Exception as e:
    test_result("DOCX export functionality", False, str(e))

# Test 5: Documentation completeness
print("\n[5] Testing documentation...")
try:
    readme_path = BASE_DIR / "docs/piyut-talis-agent/README.md"
    usage_path = BASE_DIR / "docs/piyut-talis-agent/USAGE.md"
    
    readme_exists = readme_path.exists()
    usage_exists = usage_path.exists()
    
    test_result(
        "README.md exists",
        readme_exists,
        f"Size: {readme_path.stat().st_size if readme_exists else 0} bytes"
    )
    
    test_result(
        "USAGE.md exists",
        usage_exists,
        f"Size: {usage_path.stat().st_size if usage_exists else 0} bytes"
    )
except Exception as e:
    test_result("Documentation completeness", False, str(e))

# Test 6: Cross-references
print("\n[6] Testing cross-references...")
try:
    # Check if agent spec references unified seed
    if 'agent_spec' in locals():
        integration = agent_spec.get('integration', {})
        refs_seed = 'unified_seed' in integration
        test_result(
            "Agent spec references unified seed",
            refs_seed,
            f"Path: {integration.get('unified_seed', 'N/A')}"
        )
    
    # Check if unified seed references agent spec
    if 'unified_seed' in locals() and 'agents' in unified_seed:
        piyut_agent = unified_seed['agents'].get('piyut_talis', {})
        refs_spec = 'spec_file' in piyut_agent
        test_result(
            "Unified seed references agent spec",
            refs_spec,
            f"Path: {piyut_agent.get('spec_file', 'N/A')}"
        )
except Exception as e:
    test_result("Cross-references", False, str(e))

# Summary
print("\n" + "=" * 60)
print("Integration Test Summary")
print("=" * 60)

passed = sum(1 for _, p, _ in results if p)
total = len(results)
success_rate = (passed / total * 100) if total > 0 else 0

print(f"\nTotal tests: {total}")
print(f"Passed: {passed}")
print(f"Failed: {total - passed}")
print(f"Success rate: {success_rate:.1f}%")

if passed == total:
    print("\n✓ All integration tests passed!")
    sys.exit(0)
else:
    print(f"\n✗ {total - passed} test(s) failed")
    sys.exit(1)
