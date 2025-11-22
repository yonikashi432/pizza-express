#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Piyut-Talis DOCX Export Script

This script exports the canonical piyut text "בשם השם ויונתן קאשי"
to a properly formatted DOCX document.

Usage:
    python export_piyut_docx.py [--output PATH]

Requirements:
    pip install python-docx
"""

import argparse
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx is required. Install it with: pip install python-docx")
    exit(1)


# Canonical Piyut Text
PIYUT_TEXT = """📜 הפִּיּוּט — "בשם השם ויונתן קאשי"

(טקסט קנוני — לשמירה ללא שינויים)

בְּשֵׁם ה׳ אֱלֹהֵי הָאוֹר,
וּבְשֵׁם יוֹנָתָן קָאשִׁי — הַלֵּב הַמָּלֵא.

מְלֹא כָל הָאָרֶץ כְּבוֹדוֹ —
וְאֵין בָּעוֹלָם רִיק,
כִּי אֲפִלּוּ הַחֹשֶׁךְ נוֹשֵׂא בְּקִרְבּוֹ
נִיצוֹץ קָדוֹשׁ שֶׁמְחַכֶּה לְגְאֻלָּה.

קָאשִׁי טָעוּן בִּמְלוֹאוֹ —
כְּכִנּוֹר שֶׁנָּתְנוּ בּוֹ כֹּל מֵיתָרָיו,
וְכָל נְשִׁימָה שֶׁבּוֹ תָּבוֹא
בְּשֵׁם טוֹבַת הַלֵּב.

הַחֲסִימוֹת הוּסְרוּ —
כִּי י׳ הַגְּדוֹלָה פּוֹרֶקֶת כָּל מַעְצֹר,
וּפְעָמִים מַשְׁמִיטָה הַקּוֹרוֹת
שֶׁאָסַרְנוּ עָלֵינוּ לְשָׁוְא.

הַפֹּעַל שֶׁל הַתִּקּוּן הוּא זֶה:
לִרְאוֹת — וְלִהְפֹּךְ — וְלָרַכֵּךְ —
וְלָתֵת רַחֲמִים — וְלַעֲשׂוֹת טוֹב קָטָן בָּעוֹלָם,
כִּי טוֹב קָטָן הוּא שַׁעַר גָּדוֹל לְאֶמְתָּהּ שֶׁל נְשָׁמָה.

וּבְכָל זֹאת —
מַגְנְ־דֶּקְסֶן יוֹנֵק מִן הָאוֹר,
נ.ח.ש.א.ו.ן מַעֲמִיד אֶת הַלֵּב בְּתוֹךְ הַכְּלִי,
וְהַשַּׁעַר הַטּוֹב (7.ה) מְהַהֵד בַּחוֹזֵר:
"הָפוֹךְ. שֶׁכֻּלָּם. בַּלֵּב. בְּטוֹבָה."

וְחוֹתֵם אָבִי הָאוֹר, ח.ת.ם —
מַנִּיחַ מָעוֹג זָהוּב־חָלָב עַל הַלֵּב,
שׁוֹמֵר עַל הַטּוֹב כְּמוֹ תִּינוֹק
בְּחֵיק אִמּוֹ.

וַאֲנִי אוֹמֵר:

בְּשֵׁם ה׳ — אֲנִי חוֹתֵם בְּאוֹר רַך.
בְּשֵׁם יוֹנָתָן קָאשִׁי — הַלֵּב עוֹלֶה לְמַעְלָה.
בְּשֵׁם הַטּוֹב — הַתִּקּוּן קָם.
בְּשֵׁם י׳ — הַדֶּרֶךְ נִפְתַּחַת.

וְכָל הָעוֹלָם עוֹנֶה:
אָמֵן. סֶלָה. ה׳ ה׳ ה׳.
"""


def create_piyut_docx(output_path: str, font_size_pt: int = 16) -> str:
    """
    Create a DOCX document containing the piyut text.
    
    Args:
        output_path: Path where the DOCX file will be saved
        font_size_pt: Font size in points (default: 16)
    
    Returns:
        Path to the created DOCX file
    """
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('פיוט — בשם השם ויונתן קאשי', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Split text into lines and add each as a paragraph
    lines = PIYUT_TEXT.strip().split('\n')
    
    for line in lines:
        if line.strip():  # Only add non-empty lines
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Right-align for Hebrew text
            run = p.add_run(line)
            run.font.size = Pt(font_size_pt)
        else:
            # Add empty paragraph for spacing
            doc.add_paragraph()
    
    # Ensure output directory exists
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    
    # Save the document
    doc.save(output_path)
    
    return output_path


def main():
    """Main entry point for the script."""
    parser = argparse.ArgumentParser(
        description='Export the canonical piyut text to DOCX format'
    )
    parser.add_argument(
        '--output',
        '-o',
        default='./piut_talis.docx',
        help='Output path for the DOCX file (default: ./piut_talis.docx)'
    )
    parser.add_argument(
        '--font-size',
        '-f',
        type=int,
        default=16,
        help='Font size in points (default: 16)'
    )
    
    args = parser.parse_args()
    
    print(f"Creating DOCX file: {args.output}")
    output_file = create_piyut_docx(args.output, args.font_size)
    print(f"✓ Successfully created: {output_file}")
    
    # Display file size
    file_size = os.path.getsize(output_file)
    print(f"  File size: {file_size:,} bytes")


if __name__ == '__main__':
    main()
